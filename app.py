from flask import Flask, request, jsonify, send_file
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from io import BytesIO
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import json
import re

app = Flask(__name__)

# 分解并获取项目信息
def parse_project(project_str):
    """解析单个项目字符串"""
    # 分割字段并清理空白
    raw_fields = [f.strip() for f in project_str.split("|")]
    # 去除首尾空字段（原始格式为 |字段1|字段2|...| ）
    processed_fields = raw_fields[1:-1]
    
    # 验证字段数量（根据示例应为11个字段）
    if len(processed_fields) != 11:
        print("字段数量异常: {len(processed_fields)}")
        print("异常数据:", processed_fields)
        return None
    
    # 特殊处理关键节点字段（示例中第4个字段，索引3）
    # 将分号分隔转换为斜杠分隔
    processed_fields[4] = processed_fields[4].replace('；', ';')  # 统一分隔符
    processed_fields[4] = '/'.join([x.split('：')[-1].strip() for x in processed_fields[4].split(';')])
    
    return processed_fields

@app.route('/smart', methods=['POST'])
def export_projects():
    try:
        # 增强的请求验证
        if not request.is_json:
            return jsonify({"status": "error", "message": "Content-Type必须是application/json"}), 400

        data = request.get_json(silent=True)
        
        # 验证1: 是否是有效JSON
        if data is None:  # 注意：空列表 [] 会通过此检查，但 isinstance([], dict) 会失败
            return jsonify({"status": "error", "message": "无效的JSON格式"}), 400

        # 创建Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "项目清单"
        
        # 设置表头（根据字段顺序调整）
        headers = [
            "编号", "项目名称", "负责人", "工作内容",
            "成果指标", "开始时间", "结束时间", "总预算(万)",
            "Smart符合性", "改进建议", "预算审核"
        ]
        ws.append(headers)
        ws.row_dimensions[1].height = 40  # 调整为合适值

        # ========== 关键修复3：设置列宽 ==========
        column_widths = {
            'A': 8,   # 编号
            'B': 15,  # 项目名称
            'C': 10,  # 负责人
            'D': 40,  # 工作内容
            'E': 30,  # 成果指标
            'F': 12,  # 开始时间
            'G': 12,  # 结束时间
            'H': 12,  # 熊预算
            'I': 12,  # SMART符合性
            'J': 50,  # SMART改进建议
            'K': 40   # 预算审核
        }

        # 处理数据
        for raw_item in data:
            fields = parse_project(raw_item)
            if fields:
                # 调整字段顺序（如果需要）
                # 转换预算值为数字类型
                try:
                    fields[7] = float(fields[7])  # 预算(万)转为数字
                except ValueError:
                    fields[7] = 0.0
                
                # 写入Excel
                ws.append([
                    fields[0],  # 序号
                    fields[1],  # 项目名称
                    fields[2],  # 负责人
                    fields[3],  # 工作内容
                    fields[4],  # 关键节点（已处理）
                    fields[5],  # 开始时间
                    fields[6],  # 结束时间
                    fields[7],  # 预算(万)
                    fields[8],  # 符合情况
                    fields[9],  # 存在问题
                    fields[10]  # 预算建议
                ])

        # 定义表头填充颜色（浅蓝色）
        header_fill = PatternFill(
            start_color="87CEEB",  # RGB颜色代码（浅蓝色）
            end_color="87CEEB",
            fill_type="solid"       # 纯色填充
        )

        header_fill2 = PatternFill(
            start_color="BACEEB",  # RGB颜色代码（浅蓝色）
            end_color="BACEEB",
            fill_type="solid"       # 纯色填充
        )

        # 2. 定义表头字体样式（加粗、白色字体）
        header_font = Font(
            name="微软雅黑",
            size=14,
            bold=True,
            color="FFFFFF"          # 白色字体
        )

        # 定义样式
        alignment_left = Alignment(wrap_text=True, vertical="center", horizontal="left")
        alignment_center = Alignment(wrap_text=True, vertical="center", horizontal="center")
        content_font = Font(name="微软雅黑", size=10, bold=False, color="000000")
        # 设置列宽
        for col, width in column_widths.items():
            ws.column_dimensions[col].width = width

        # 遍历所有单元格设置样式
        for row in ws.iter_rows():  # 遍历所有行
            for cell in row:        # 遍历行内所有单元格
                cell.font = content_font
                # 根据列字母设置对齐方式
                if cell.column_letter in ['D', 'E', 'J', 'K']:
                    cell.alignment = alignment_left
                else:
                    cell.alignment = alignment_center

        for col_idx in range(1, len(headers)+1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = header_font
            cell.alignment = alignment_center
            # 按列位置设置不同背景色
            if col_idx > 8:  # 从第I列(索引9)开始
                cell.fill = header_fill2
            else:
                cell.fill = header_fill

        # 保存到内存缓冲区
        excel_buffer = BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)

        # 返回Excel文件
        return send_file(
            excel_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='项目信息SMART评估后.xlsx'
        )

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"处理失败: {str(e)}"
        }), 500

def parse_records(line):
    parts = [p.strip() for p in line.split('|') if p.strip()]
    print("parts=") 
    print(parts)
    if len(parts) != 14:
        return record
    # 类型转换处理
    try:
        record = {
            1:parts[0],   # 编号
            2:parts[1],  # 项目名称
            3:parts[2],  # 负责人
            4:parts[3],  # 工作内容
            5:parts[4],  # 成果指标
            6:parts[5],  # 开始时间
            7:parts[6],  # 结束时间
            8:parts[7],  # 总预算
            9:parts[8], # 条目编号
            10:parts[9], # 条目描述
            11:parts[10], # 交付物
            12:parts[11], # 交付时间）
            13:parts[12], # 工作量估计
            14:parts[13]  # 成本估计
        }
        print("record=")
        print(record)
    except (ValueError, IndexError) as e:
        print(f"数据解析错误：{line} -> {str(e)}")
    
    return record

@app.route('/detail', methods=['POST'])
def detail_projects():
    try:
        # 增强的请求验证
        if not request.is_json:
            return jsonify({"status": "error", "message": "Content-Type必须是application/json"}), 400

        data = request.get_json(silent=True)

        # 验证1: 是否是有效JSON
        if data is None:  # 注意：空列表 [] 会通过此检查，但 isinstance([], dict) 会失败
            return jsonify({"status": "error", "message": "无效的JSON格式"}), 400

        # 创建Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "项目分解"
        
        # 设置表头（根据字段顺序调整）
        headers = ["项目编号", "项目名称", "负责人", "工作内容", "成果指标", "开始时间", "结束时间", "总预算(万)",
                   "条目编号", "条目描述", "条目交付物", "条目交付时间", "条目工作量估计", "条目成本估计"]
        ws.append(headers)
        ws.row_dimensions[1].height = 40  # 调整为合适值

        # 定义表头填充颜色（浅蓝色）
        header_fill = PatternFill(
            start_color="87CEEB",  # RGB颜色代码（浅蓝色）
            end_color="87CEEB",
            fill_type="solid"       # 纯色填充
        )

        header_fill2 = PatternFill(
            start_color="BACEEB",  # RGB颜色代码（浅蓝色）
            end_color="BACEEB",
            fill_type="solid"       # 纯色填充
        )

        # 2. 定义表头字体样式（加粗、白色字体）
        header_font = Font(
            name="微软雅黑",
            size=14,
            bold=True,
            color="FFFFFF"          # 白色字体
        )

        # 定义样式
        alignment_left = Alignment(wrap_text=True, vertical="center", horizontal="left")
        alignment_center = Alignment(wrap_text=True, vertical="center", horizontal="center")
        content_font = Font(name="微软雅黑", size=10, bold=False, color="000000")

        print("set table title=")
        for col_idx in range(1, len(headers)+1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = header_font
            cell.alignment = alignment_center
            # 按列位置设置不同背景色
            if col_idx > 8:  # 从第I列(索引9)开始
                cell.fill = header_fill
            else:
                cell.fill = header_fill2
        # ========== 关键修复3：设置列宽 ==========
        column_widths = {
            'A': 8,   # 项目编号
            'B': 15,  # 项目名称
            'C': 10,  # 负责人
            'D': 40,  # 工作内容
            'E': 30,  # 成果指标
            'F': 12,  # 开始时间
            'G': 12,  # 结束时间
            'H': 12,  # 总预算
            'I': 8,   # 条目编号
            'J': 25,  # 条目描述
            'K': 25,  # 条目交付物
            'L': 12,  # 条目交付时间
            'M': 12,  # 工作量估计
            'N': 12   # 成本估计
        }
        print("set kuandu=")
        # 设置列宽
        for col, width in column_widths.items():
            ws.column_dimensions[col].width = width

        # 处理数据
        start_row = 1
        end_row = 1
        for raw_item in data:
            lines = raw_item.strip().split('\n')
            print("lines=")
            print(lines)
            start_row = end_row
            for line in lines:
                print("line=")
                print(line)
                fields = parse_records(line)
                max_index = max(fields.keys(), default=0)
                rec = [None] * max_index
                for index, recValue in fields.items():
                    rec[index - 1] = recValue

                print("fields=")
                print(fields)
                if len(fields) == 14:
                    # 写入Excel
                    ws.append([
                        rec[0],  # 项目编号
                        rec[1],  # 项目名称
                        rec[2],  # 负责人
                        rec[3],  # 工作内容
                        rec[4],  # 成果指标
                        rec[5],  # 开始时间
                        rec[6],  # 结束时间
                        rec[7],  # 总预算(万)
                        rec[8],  # 条目编号
                        rec[9],  # 条目描述
                        rec[10],  # 条目交付物
                        rec[11],  # 条目交付时间
                        rec[12],  # 工作量估计
                        rec[13]   # 成本估计
                    ])
                end_row = end_row + 1

                print("set style=")
                # 遍历所有单元格设置样式
                for row in ws.iter_rows(min_row=end_row, max_row=end_row):  # 遍历当前行
                    for cell in row:        # 遍历行内所有单元格
                        print("set style 1111=")
                        cell.font = content_font
                        print("set style 222=")
                        # 根据列字母设置对齐方式
                        if cell.column_letter in ['D', 'E', "J", "K"]:
                            cell.alignment = alignment_left
                        else:
                            cell.alignment = alignment_center
            if(end_row > start_row + 1):
                ws.merge_cells(start_row=start_row+1, start_column=1, end_row=end_row, end_column=1)
                ws.merge_cells(start_row=start_row+1, start_column=2, end_row=end_row, end_column=2)
                ws.merge_cells(start_row=start_row+1, start_column=3, end_row=end_row, end_column=3)
                ws.merge_cells(start_row=start_row+1, start_column=4, end_row=end_row, end_column=4)
                ws.merge_cells(start_row=start_row+1, start_column=5, end_row=end_row, end_column=5)
                ws.merge_cells(start_row=start_row+1, start_column=6, end_row=end_row, end_column=6)
                ws.merge_cells(start_row=start_row+1, start_column=7, end_row=end_row, end_column=7)
                ws.merge_cells(start_row=start_row+1, start_column=8, end_row=end_row, end_column=8)

        # 保存到内存缓冲区
        excel_buffer = BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)

        # 返回Excel文件
        return send_file(
            excel_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='项目信息分解建议.xlsx'
        )

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"处理失败: {str(e)}"
        }), 500

def create_word_document(data):
    """创建Word文档并保存"""
    try:
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 添加文档标题
        title = doc.add_heading(level=0)
        title_run = title.add_run(f"{data['Document_Framework']['title']} (版本：{data['document_structure']['version']})")
        title_run.font.size = Pt(16)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 处理主要章节
        for section in data['document_structure']['sections']:
            # 添加章节标题
            heading = doc.add_heading(section['section_title'], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # 处理子章节
            for subsection in section.get('subsections', []):
                # 添加子章节标题
                sub_heading = doc.add_heading(subsection['subsection_title'], level=2)
                
                # 添加内容段落
                content = doc.add_paragraph()
                content.add_run(subsection['content']).bold = False
                content.paragraph_format.space_after = Pt(6)
        
        # 添加知识引用章节
        doc.add_heading("知识引用依据", level=1)
        for source in data['knowledge_retrieval_results']['sources']:
            doc.add_heading(f"来源：{source['source']}", level=2)
            
            # 添加关键点列表
            for point in source['key_points']:
                p = doc.add_paragraph(style='ListBullet')
                p.add_run(point)

        # 保存到内存缓冲区
        excel_buffer = BytesIO()
        doc.save(excel_buffer)
        excel_buffer.seek(0)

        # 返回Excel文件
        return send_file(
            excel_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='需求说明书框架.xlsx'
        )

        return True
    except Exception as e:
        raise Exception(f"文档生成失败: {str(e)}")

@app.route('/convert', methods=['POST'])
def process_requirements():
    """完整处理流程"""
    try:
        # 获取原始JSON数据
        response = request.get_json()
        response = request.get_data()
        # 验证1: 是否是有效JSON
        if response is None:  # 注意：空列表 [] 会通过此检查，但 isinstance([], dict) 会失败
            return jsonify({"status": "error", "message": "无效的JSON格式"}), 400

        # 解析JSON数据
        data = json.loads(response)
        
        # 生成Word文档
        if create_word_document(data,):
            return f"文档已成功生成"
        return "文档生成失败"
    except json.JSONDecodeError as e:
        return f"JSON解析错误：{str(e)}"
    except Exception as e:
        return f"处理过程中发生错误：{str(e)}"


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5996, debug=True)
